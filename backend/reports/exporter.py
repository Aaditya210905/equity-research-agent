import io
from docx import Document
from fpdf import FPDF
from typing import Dict, Any

from reports.formatter import report_to_markdown, report_to_html

class PDFReport(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'Equity Research Report', 0, 1, 'C')

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def generate_pdf(report_data: Dict[str, Any]) -> io.BytesIO:
    """Generates a PDF using fpdf2."""
    pdf = PDFReport()
    pdf.add_page()
    pdf.set_font("Arial", size=11)
    
    metadata = report_data.get("metadata", {})
    sections = report_data.get("sections", [])
    
    # Title
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, f"{metadata.get('company', 'Company')} ({metadata.get('ticker', 'TICKER')})", ln=True)
    pdf.set_font("Arial", size=10)
    pdf.cell(0, 10, f"Generated At: {metadata.get('generated_at', '')}", ln=True)
    pdf.ln(5)
    
    for section in sections:
        title = section.get("title", "")
        content = section.get("content", "")
        
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 10, title, ln=True)
        
        pdf.set_font("Arial", size=11)
        if content:
            # MultiCell handles line breaks
            pdf.multi_cell(0, 8, content.encode('latin-1', 'replace').decode('latin-1'))
            pdf.ln(5)
            
    pdf_bytes = pdf.output(dest='S')
    return io.BytesIO(pdf_bytes)

def generate_docx(report_data: Dict[str, Any]) -> io.BytesIO:
    """Generates a DOCX document using python-docx."""
    doc = Document()
    
    metadata = report_data.get("metadata", {})
    sections = report_data.get("sections", [])
    
    doc.add_heading(f"Equity Research Report: {metadata.get('company', 'Company')} ({metadata.get('ticker', 'TICKER')})", 0)
    doc.add_paragraph(f"Generated At: {metadata.get('generated_at', '')}")
    
    for section in sections:
        title = section.get("title", "")
        content = section.get("content", "")
        
        doc.add_heading(title, level=1)
        if content:
            doc.add_paragraph(content)
            
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io
